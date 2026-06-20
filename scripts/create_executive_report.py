"""Create the final executive business report DOCX.

The report consolidates Phase 1-6 outputs into a 2-3 page leadership brief.
It is intentionally concise: the audience is hospital leadership deciding
whether to sponsor a controlled deployment and operational adoption.
"""

from __future__ import annotations

import csv
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOCX = PROJECT_ROOT / "Healthcare_Insights_Report.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(95, 105, 119)
GRAY_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
WHITE = RGBColor(255, 255, 255)
BLACK = RGBColor(0, 0, 0)


def load_json(relative_path: str) -> dict[str, Any]:
    return json.loads((PROJECT_ROOT / relative_path).read_text(encoding="utf-8"))


def load_csv(relative_path: str) -> list[dict[str, str]]:
    with (PROJECT_ROOT / relative_path).open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def money(value: float) -> str:
    return f"${value / 1_000_000:.1f}M"


def pct(value: float) -> str:
    return f"{value * 100:.1f}%"


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def paragraph_border_bottom(paragraph, color: str = "D7DBE2", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    existing_grid = tbl.find(qn("w:tblGrid"))
    if existing_grid is not None:
        tbl.remove(existing_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths_dxa[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            set_cell_margins(cell)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 9.2, color: RGBColor = BLACK) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    run = paragraph.add_run(text)
    set_run_font(run, size=size, color=color, bold=bold)


def add_body_paragraph(doc: Document, text: str, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.10
    if bold_lead:
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, size=11, color=INK, bold=True)
        rest = paragraph.add_run(text)
    else:
        rest = paragraph.add_run(text)
    set_run_font(rest, size=11, color=BLACK)


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=BLACK)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.167
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, color=BLACK)


def add_callout(doc: Document, label: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.10
    lead = paragraph.add_run(f"{label}: ")
    set_run_font(lead, size=11.2, color=DARK_BLUE, bold=True)
    body = paragraph.add_run(text)
    set_run_font(body, size=10.8, color=BLACK)


def add_metric_strip(doc: Document, metrics: list[tuple[str, str]]) -> None:
    widths = [1872] * 5
    table = doc.add_table(rows=2, cols=5)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for col, (label, value) in enumerate(metrics):
        label_cell = table.cell(0, col)
        value_cell = table.cell(1, col)
        set_cell_shading(label_cell, BLUE_FILL)
        set_cell_text(label_cell, label, bold=True, size=8.4, color=DARK_BLUE)
        set_cell_text(value_cell, value, bold=True, size=10.2, color=INK)
        value_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_key_findings_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [2100, 3820, 3440])
    headers = ["Leadership Area", "Evidence From Analytics", "Business Implication"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, GRAY_FILL)
        set_cell_text(cell, header, bold=True, size=9, color=DARK_BLUE)
    for area, evidence, implication in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], area, bold=True, size=9.2, color=INK)
        set_cell_text(cells[1], evidence, size=8.8)
        set_cell_text(cells[2], implication, size=8.8)


def add_model_table(doc: Document, rows: list[tuple[str, str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [1750, 2050, 2020, 3540])
    headers = ["Model", "Selected Method", "Business Recall", "Executive Interpretation"]
    for idx, header in enumerate(headers):
        cell = table.cell(0, idx)
        set_cell_shading(cell, GRAY_FILL)
        set_cell_text(cell, header, bold=True, size=9, color=DARK_BLUE)
    for model, method, recall, interpretation in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], model, bold=True, size=9.1, color=INK)
        set_cell_text(cells[1], method, size=8.8)
        set_cell_text(cells[2], recall, bold=True, size=8.8, color=INK)
        set_cell_text(cells[3], interpretation, size=8.8)


def add_architecture_table(doc: Document) -> None:
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    set_table_geometry(table, [1560, 1560, 1560, 1560, 1560, 1560])
    stages = [
        ("Raw data", "patients, visits, billing CSVs"),
        ("SQL layer", "integrity, indexes, reusable views"),
        ("Feature table", "25,000 rows, 69 columns"),
        ("Models", "risk and claim classifiers"),
        ("API", "FastAPI scoring and logging"),
        ("Governance", "validation, drift, retraining"),
    ]
    for idx, (stage, detail) in enumerate(stages):
        cell = table.cell(0, idx)
        set_cell_shading(cell, BLUE_FILL if idx % 2 == 0 else GRAY_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        title = paragraph.add_run(stage)
        set_run_font(title, size=9, color=DARK_BLUE, bold=True)
        paragraph.add_run("\n")
        detail_run = paragraph.add_run(detail)
        set_run_font(detail_run, size=7.7, color=BLACK)


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Healthcare AI Capstone | Executive Business Report")
    set_run_font(run, size=8.5, color=MUTED)
    paragraph_border_bottom(paragraph, color="D7DBE2", size="4")

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Decision-support prototype. Use with human review, monitoring, and governance controls.")
    set_run_font(run, size=8, color=MUTED)


def add_masthead(doc: Document) -> None:
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.line_spacing = 1.0
    title_run = title.add_run("Healthcare Insights Report")
    set_run_font(title_run, size=23, color=INK, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    subtitle_run = subtitle.add_run("Hospital Operations & Revenue Risk Intelligence Platform")
    set_run_font(subtitle_run, size=13, color=MUTED)

    metadata = [
        ("Audience", "Hospital leadership and capstone evaluation panel"),
        ("Decision focus", "Sponsor controlled deployment and operational adoption"),
        ("Data foundation", "25,000 integrated encounters from patients, visits, and billing"),
        ("Prepared", "June 20, 2026"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(f"{label}: ")
        set_run_font(label_run, size=10.5, color=INK, bold=True)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=BLACK)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_after = Pt(8)
    paragraph_border_bottom(rule, color="2E74B5", size="10")


def build_report() -> None:
    feature_schema = load_json("data_outputs/feature_schema.json")
    profile = feature_schema["profile"]
    model_card = load_json("data_outputs/phase4/model_card.json")
    phase5 = load_json("data_outputs/phase5/phase5_verification_summary.json")
    phase6 = load_json("data_outputs/phase6/drift_detection_summary.json")
    rejection_rows = load_csv("data_outputs/phase1/financial_top5_insurance_providers_by_rejection_rate.csv")
    los_rows = load_csv("data_outputs/phase1/operational_top5_departments_by_avg_los.csv")
    high_risk_rows = load_csv("data_outputs/phase1/operational_high_risk_pct_by_department.csv")
    volume_rows = load_csv("data_outputs/phase1/operational_top10_departments_by_visit_volume.csv")

    total_rows = profile["model_table_rows"]
    high_risk_count = profile["risk_score_distribution"]["High"]
    rejected_count = profile["claim_status_distribution"]["Rejected"]
    revenue_gap = profile["total_billed_amount"] - profile["total_approved_amount"]
    top_volume = volume_rows[0]
    top_los = los_rows[0]
    top_high_risk = high_risk_rows[0]
    top_rejection = rejection_rows[0]

    risk_metrics = model_card["models"]["risk"]["test_metrics"]
    claim_metrics = model_card["models"]["claim"]["test_metrics"]

    doc = Document()
    set_document_styles(doc)
    add_header_footer(doc)
    add_masthead(doc)

    add_callout(
        doc,
        "Executive decision",
        "Proceed with a controlled pilot, not full automation. The platform is ready to support leadership dashboards and review queues, but current model performance and data-quality findings require human oversight, remediation, and monitored retraining before policy automation.",
    )

    add_heading(doc, "Portfolio Snapshot", 1)
    add_metric_strip(
        doc,
        [
            ("Encounters", f"{total_rows:,}"),
            ("Billed amount", money(profile["total_billed_amount"])),
            ("Approved amount", money(profile["total_approved_amount"])),
            ("Realization ratio", pct(profile["overall_revenue_realization_ratio"])),
            ("Revenue gap", money(revenue_gap)),
        ],
    )

    add_heading(doc, "Key Operational and Financial Findings", 1)
    add_key_findings_table(
        doc,
        [
            (
                "Patient flow",
                f"{top_volume['department']} carried the highest visit volume ({int(top_volume['total_visits']):,}); {top_los['department']} had the highest average length of stay ({float(top_los['avg_length_of_stay_hours']):.2f} hours).",
                "Operational dashboards should monitor department volume and length-of-stay together, because capacity pressure is not only a volume issue.",
            ),
            (
                "Risk concentration",
                f"High Risk visits represented {high_risk_count:,} of {total_rows:,} encounters ({high_risk_count / total_rows:.1%}); {top_high_risk['department']} had the highest High Risk share ({float(top_high_risk['high_risk_visit_pct']):.2f}%).",
                "ER and ICU triage should be prioritized for pilot workflows, staffing review, and high-risk escalation design.",
            ),
            (
                "Revenue leakage",
                f"The portfolio billed {money(profile['total_billed_amount'])} and approved {money(profile['total_approved_amount'])}, leaving a {money(revenue_gap)} gap; {top_rejection['insurance_provider']} had the highest rejection rate ({float(top_rejection['claim_rejection_rate_pct']):.2f}%).",
                "Finance teams should use payer-specific review queues before submission, focused first on high-bill and high-rejection patterns.",
            ),
            (
                "Data reliability",
                f"Approved amount is missing in {profile['missing_approved_amount_count']:,} records, payment days in {profile['missing_payment_days_count']:,}, and {profile['high_billed_zero_or_missing_approved_count']:,} high-billed records have zero or missing approvals.",
                "Leadership reporting should include data-quality gates so revenue and payment-delay decisions are not made on incomplete claim outcomes.",
            ),
        ],
    )

    add_heading(doc, "Model Impact Summary", 1)
    add_body_paragraph(
        doc,
        "The models create a practical triage layer: they can prioritize work for staff review, but they should not replace clinical judgment or billing policy. Phase 4 evaluation showed modest predictive signal, which is useful for early review queues but not sufficient for autonomous decisions.",
    )
    add_model_table(
        doc,
        [
            (
                "Visit risk",
                "Logistic Regression",
                f"High Risk recall {risk_metrics['business_recall']:.1%}; macro F1 {risk_metrics['macro_f1']:.3f}",
                "Use as an early warning signal for operational review, especially in ER/ICU workflows; continue feature enrichment before clinical policy use.",
            ),
            (
                "Claim outcome",
                "Random Forest",
                f"Rejected recall {claim_metrics['business_recall']:.1%}; macro F1 {claim_metrics['macro_f1']:.3f}",
                "Use to prioritize pre-submission claim review and payer follow-up; it is stronger as revenue-cycle decision support than as automatic disposition.",
            ),
        ],
    )

    add_heading(doc, "Deployment, Architecture, and Governance", 1)
    add_architecture_table(doc)
    add_body_paragraph(
        doc,
        f"Phase 5 operationalized the selected models through FastAPI endpoints for health, version metadata, risk prediction, and claim prediction. Verification confirmed live scoring and prediction logging, including model version and feature hash metadata. Phase 6 added validation, feature drift, prediction drift, and audit-log monitoring.",
    )
    add_bullet(doc, f"API readiness: verified endpoints include {', '.join(phase5['verified_endpoints'])}; prediction logging is verified at {phase5['audit_log_path']}.")
    add_bullet(doc, f"Monitoring status: {phase6['data_validation']['checks']} validation checks ran, with {phase6['data_validation']['failed_checks']} failures and {phase6['data_validation']['warning_checks']} warnings.")
    add_bullet(doc, f"Drift status: risk prediction drift is stable (PSI {phase6['prediction_drift']['tasks'][0]['prediction_psi']:.3f}); claim prediction drift is watch level (PSI {phase6['prediction_drift']['tasks'][1]['prediction_psi']:.3f}).")

    add_heading(doc, "Business Recommendations for Hospital Leadership", 1)
    add_numbered(doc, "Launch a 60-90 day decision-support pilot in ER/ICU operations and revenue-cycle claim review, with staff feedback captured against every High Risk or Rejected alert.")
    add_numbered(doc, "Create a data-quality remediation workstream for approved amount, payment days, temporal anomalies, and high-billed zero/missing approval cases before expanding executive financial dashboards.")
    add_numbered(doc, "Use payer-specific revenue optimization: start with SecureLife and MediCareX rejection review, then add rules for high-billed claims and payer response delays.")
    add_numbered(doc, "Govern the platform as a monitored AI product: retain model cards, audit logs, drift reports, fairness checks, and retraining approvals before promoting any new model version.")

    add_heading(doc, "Strategic Value", 1)
    add_body_paragraph(
        doc,
        "The capstone has converted raw hospital operations and billing data into a governed analytics product: a SQL foundation, modeling-ready feature layer, two classification systems, an API deployment path, and monitoring controls. The immediate leadership value is better visibility into patient-flow pressure and revenue leakage risk. The next value unlock is disciplined pilot adoption, not unchecked automation.",
    )

    doc.core_properties.title = "Healthcare Insights Report"
    doc.core_properties.subject = "Executive business presentation for hospital operations and revenue risk intelligence"
    doc.core_properties.author = "Healthcare AI Capstone Project"
    doc.save(OUTPUT_DOCX)
    print(f"Created {OUTPUT_DOCX}")


if __name__ == "__main__":
    build_report()
